from django.shortcuts import render,render_to_response
from django.http import HttpResponse
from django.template import loader
from django.core.files.storage import FileSystemStorage
from . forms import CurrentTweetForm,CurrentTweetHashtagForm,pdf_lda_mod,pdf_lda_mode,EmailForm,lda_mod,csv_lda_mod,text_lda_mod,text_lda_mode,OldTweetForm,GetUserTimelineForm,GetFollowerslistForm
import os
import csv
from django.views.decorators.csrf import csrf_protect
from django.core.mail import send_mail, EmailMessage
from django.core.mail import mail_admins
import tweepy
from posts.utils import *
import pandas as pd
import GetOldTweets3 as got3
from datetime import date
import speech_recognition as sr
from docx import Document
from docx.shared import Inches
from textblob import TextBlob
import sys
import io
import numpy as np
from datetime import datetime
import re
import nltk
from nltk.corpus import wordnet
from nltk.corpus import stopwords
from nltk.stem import WordNetLemmatizer
from nltk.stem.porter import PorterStemmer
from nltk.corpus import stopwords 
from nltk.stem.wordnet import WordNetLemmatizer
import string
import gensim
from gensim import corpora
import PyPDF2


from collections import Counter
#somewhere import handle_uploaded_file

#from .models import ModelWithFileField

from django.contrib.auth.forms import UserCreationForm,AuthenticationForm


from django.http import HttpResponse

# Create your views here.

def home(request):
	#return HttpResponse("Hello My Love")
	#template=loader.get_template('posts/home.html')
	return render(request,"home.html",{})

def twitterdata(request):
	return render(request,'twitterdata.html',{})

def topicmodeling(request):
	return render(request,'topicmodeling.html',{})

def speechrecognition(request):
	return render(request,'speech.html',{})

def textdata(request):
	return render(request,'upload.html',{})

def csvdata(request):
	return render(request, 'uploadcsv.html',{})

def sentiment(request):
	return render(request,'sentiment.html',{})
	
#def textdata(request):
	#return render(request,'upload.html',{})

def currenttweetsform(request):
	return render(request,'currenttweetsform.html',{})

def oldtweets(request):
	return render(request,"oldtweets.html",{})

def getusertimeline(request):
	return render(request,"getusertimeline.html",{})

def pdfdata(request):
	return render(request,"pdflda.html",{})

#def upload(request):
	#return render(request,"upload.html",{})


#def currenttweetsform(request):
	#if request.method=='POST':
		#form=CurrentTweetForm(request.POST)
		#if form.is_valid():
			#username=form.cleaned_data['username']
			#items=form.cleaned_data['items']
			#username = request.POST.get('username')
			#items = request.POST.get('items')
			
			#auth = tweepy.OAuthHandler(CONSUMER_KEY, CONSUMER_SECRET)
			#auth.set_access_token(ACCESS_TOKEN, ACCESS_TOKEN_SECRET)
			#api = tweepy.API(auth,wait_on_rate_limit=True)
  
			#users= tweepy.Cursor(api.user_timeline,screen_name=username,count=200).items(int(items))
			#users= tweepy.Cursor(api.search, q= username,count=200).items(int(items))
			#outtweets=[]
			#outtweets=[[tweet.id_str,tweet.user.created_at,tweet.user.screen_name,tweet.user.name.encode('utf-8'),tweet.user.description.encode('utf-8'),tweet.user.statuses_count,tweet.user.followers_count,tweet.created_at,tweet.text.encode('utf-8'),len(tweet.text),tweet.source,tweet.favorite_count,tweet.retweet_count,tweet.user.location,tweet.entities.get('media', [])] for tweet in users]
			#outtweets = [[d.id,d.created_at,d.user.location,d.text] for d in users]
			#outtweets= pd.DataFrame(outtweets)
			#outtweets.columns=['id','created_at','location','text']
			#outtweets.columns=['id','user created_at','screenname','name','description','Tweets_count','follower_count','created_at','text','length','source','favorite_count','retweets','location','media_url']
			#data=outtweets.to_csv()
			#print(data)
			#response = HttpResponse(outtweets.to_csv())
			#response['Content-Disposition'] = "attachment; filename=userscurrenttweets.csv"
			#response['Content-Type'] = 'text/csv'
			#return response
	#else:
		#form=CurrentTweetForm()

	#return render(request,'currenttweetsform.html',{})

def currenttweetsform(request):
	if request.method=='POST':
		form=CurrentTweetForm(request.POST)
		if form.is_valid():
			username=form.cleaned_data['username']
			items=form.cleaned_data['items']
			recipient_email_address=form.cleaned_data["recipient_email_address"]
			#username = request.POST.get('username')
			#items = request.POST.get('items')
			
			auth = tweepy.OAuthHandler(CONSUMER_KEY, CONSUMER_SECRET)
			auth.set_access_token(ACCESS_TOKEN, ACCESS_TOKEN_SECRET)
			api = tweepy.API(auth,wait_on_rate_limit=True)
  
			#users= tweepy.Cursor(api.user_timeline,screen_name=username,count=200).items(int(items))
			users= tweepy.Cursor(api.search, q= username,count=200).items(int(items))
			#users=tweepy.Cursor(api.get_user)
			outtweets=[]

			#for tweet in users:
				#not all tweets will have media url, so lets skip them
				#try:
					#print (tweet.entities['media'][0]['media_url'])

				#except(NameError, KeyError):
						#we dont want to have any entries without the media_url so lets do nothing
					#pass
				#else:
					 #got media_url - means add it to the output
					#outtweets.append([tweet.id_str,tweet.user.screen_name,tweet.user.name.encode('utf-8'),tweet.user.description.encode('utf-8'),tweet.user.statuses_count,tweet.user.followers_count,tweet.created_at, tweet.text.encode("utf-8"),len(tweet.text),tweet.source,tweet.favorite_count,tweet.retweet_count,tweet.user.location, tweet.entities['media'][0]['media_url']])
  
			#outtweets[[tweet.id_str,tweet.user.screen_name,tweet.user.name.encode('utf-8'),tweet.user.description.encode('utf-8'),tweet.user.statuses_count,tweet.user.followers_count,tweet.created_at,tweet.text.encode('utf-8'),len(tweet.text),tweet.source,tweet.favorite_count,tweet.retweet_count,tweet.user.location,tweet.status.entities['media'] for tweet in users]
			outtweets=[[tweet.id_str,tweet.user.verified,tweet.user.created_at,tweet.user.screen_name,tweet.user.name.encode('utf-8'),tweet.user.description.encode('utf-8'),tweet.user.statuses_count,tweet.user.friends_count,tweet.user.followers_count,tweet.created_at,tweet.text.encode('utf-8'),len(tweet.text),tweet.source,tweet.favorite_count,tweet.retweet_count,tweet.user.location.encode('utf-8'),tweet.entities.get('media', [])] for tweet in users]
			outtweets= pd.DataFrame(outtweets)
			#outtweets.to_csv()
			
			outtweets.columns=['id','verified','user_created_at','screenname','name','description','Tweets_count','friends_count','follower_count','created_at','text','length','source','favorite_count','retweets','location','media_url']
			email = EmailMessage("File Downloaded", "Your file has been downloaded", "ILE@rennes-sb.com", [recipient_email_address])
			email.send(fail_silently=False)
			
			
			response = HttpResponse(outtweets.to_csv(),email)
			response['Content-Disposition'] = "attachment; filename=userscurrenttweets.csv"
			response['Content-Type'] = 'text/csv'
			return response
	else:
		form=CurrentTweetForm()

	return render(request,'currenttweetsform.html',{})


def upload_file_currenttweets(request):
	if request.method == 'POST':
		f= pd.read_csv(request.FILES["document"])
		#df=pd.DataFrame(f,columns=["username","items"])
		df=pd.DataFrame(f)
		print(df)
		data = pd.DataFrame([])
	
		for index, row in df.head().iterrows():
			auth = tweepy.OAuthHandler(CONSUMER_KEY, CONSUMER_SECRET)
			auth.set_access_token(ACCESS_TOKEN, ACCESS_TOKEN_SECRET)
			api = tweepy.API(auth,wait_on_rate_limit=True)
  
			#users= tweepy.Cursor(api.user_timeline,screen_name=row['username'],count=200).items(int(row['items']))
			users= tweepy.Cursor(api.search, q= '@' + row['username'],count=200).items(int(row['items']))
			#users= tweepy.Cursor(api.get_user,q= row["username"],count=200)
			outtweets=[]
			outtweets=[[tweet.id_str,tweet.user.verified,tweet.user.created_at,tweet.user.screen_name,tweet.user.name.encode('utf-8'),tweet.user.description.encode('utf-8'),tweet.user.statuses_count,tweet.user.friends_count,tweet.user.followers_count,tweet.created_at,tweet.text.encode('utf-8'),len(tweet.text),tweet.source,tweet.favorite_count,tweet.retweet_count,tweet.user.location.encode('utf-8'),tweet.entities.get('media', [])] for tweet in users]
			
			outtweets= pd.DataFrame(outtweets)
			#outtweets.columns=['id','created_at','location','text']
			#outtweets.columns=['id','screenname','created_at','text','length','source','favorite_count','retweets','location']
			outtweets.columns=['id','verified','user_created_at','screenname','name','description','Tweets_count','friends_count','follower_count','created_at','text','length','source','favorite_count','retweets','location','media_url']
			#outtweets.columns=['id','screenname','status_count','follower_count','friends_count','favorite_count','created_at']
			#outtweets.columns=['id','screenname','name','description','Tweets_count','follower_count','created_at','text','length','source','favorite_count','retweets','location']
			data=data.append(outtweets)


	if request.method=='POST':
		
		form = EmailForm(request.POST)
		if form.is_valid():
			recipient_email_address=form.cleaned_data["recipient_email_address"]
			email = EmailMessage("File Downloaded", "Your file has been downloaded", "ILE@rennes-sb.com", [recipient_email_address])
			email.send(fail_silently=False)		
		response = HttpResponse(data.to_csv(),email)
		response['Content-Disposition'] = "attachment; filename=userscurrenttweets.csv"
		response['Content-Type'] = 'text/csv'
		return response
	else:
		return HttpResponse("please enter proper data file")
				

		
		
	
		
	

			
	
		
	



def currenttweetshashtagform(request):
	if request.method=='POST':
		form=CurrentTweetHashtagForm(request.POST)
		if form.is_valid():
			username=form.cleaned_data['username']
			items=form.cleaned_data['items']
			recipient_email_address=form.cleaned_data["recipient_email_address"]
			auth = tweepy.OAuthHandler(CONSUMER_KEY, CONSUMER_SECRET)
			auth.set_access_token(ACCESS_TOKEN, ACCESS_TOKEN_SECRET)
			api = tweepy.API(auth,wait_on_rate_limit=True)
  
			users= tweepy.Cursor(api.search, q='#' + username,count=200).items(int(items))
			outtweets=[]
			#outtweets=[[tweet.id_str,tweet.user.screen_name,tweet.user.name,tweet.user.description.encode('utf-8'),tweet.user.statuses_count,tweet.user.followers_count,tweet.created_at, tweet.text.encode('utf-8'),len(tweet.text),tweet.source,tweet.favorite_count,tweet.retweet_count,tweet.user.location] for tweet in users]
			outtweets=[[tweet.id_str,tweet.user.verified,tweet.user.created_at,tweet.user.screen_name,tweet.user.name.encode('utf-8'),tweet.user.description.encode('utf-8'),tweet.user.statuses_count,tweet.user.friends_count,tweet.user.followers_count,tweet.created_at,tweet.text.encode('utf-8'),len(tweet.text),tweet.source,tweet.favorite_count,tweet.retweet_count,tweet.user.location.encode('utf-8'),tweet.entities.get('media', [])] for tweet in users]
			#outtweets = [[d.id,d.created_at,d.user.location,d.text] for d in users]
			outtweets= pd.DataFrame(outtweets)
			#outtweets.columns=['id','created_at','location','text']
			#outtweets.columns=['id','screenname','name','description','Tweets_count','follower_count','created_at','text','length','source','favorite_count','retweets','location']
			outtweets.columns=['id','verified','user_created_at','screenname','name','description','Tweets_count','friends_count','follower_count','created_at','text','length','source','favorite_count','retweets','location','media_url']
			email = EmailMessage("File Downloaded", "Your file has been downloaded", "ILE@rennes-sb.com", [recipient_email_address])
			email.send(fail_silently=False)
			
			response = HttpResponse(outtweets.to_csv(),email)
			response['Content-Disposition'] = "attachment; filename=usershashtagtweets.csv"
			response['Content-Type'] = 'text/csv'
			return response
	else:
		form=CurrentTweetHashtagForm()

	return render(request,'currenttweetsform.html',{})

def current_upload(request):
	upload_format = pd.DataFrame(columns=['username', 'items'])
	upload_format.set_index('username',inplace=True)
	print(('"'+"2018-12-12"+'"'))
	response = HttpResponse(upload_format.to_csv())
	response['Content-Disposition'] = "attachment; filename=upload_format.csv"
	response['Content-Type'] = 'text/csv'
	return response

def sample_upload(request):
	upload_format = pd.DataFrame(columns=['#username', 'items'])
	upload_format.set_index('#username',inplace=True)
	print(('"'+"2018-12-12"+'"'))
	response = HttpResponse(upload_format.to_csv())
	response['Content-Disposition'] = "attachment; filename=upload_format.csv"
	response['Content-Type'] = 'text/csv'
	return response

def date_upload_format(request):
	#data = {'username': ["#ronaldo","ramu","football"], 
		#'since': ['"'+"2018-12-12"+'"','"'+"2016-01-31"+'"', '"'+"2018-12-31"+'"' ],
		#'until': ['"'+"2019-09-31"+'"','"'+"2019-09-31"+'"', '"'+"2019-09-28"+'"' ], 

		#'maxtweets': [30,25,100]}
	format_file = pd.DataFrame(columns=['username', 'since','until','maxtweets'])  
	#format_file=pd.DataFrame(data)
	format_file.set_index('username',inplace=True)
	response = HttpResponse(format_file.to_csv())
	response['Content-Disposition'] = "attachment; filename=format_file.csv"
	response['Content-Type'] = 'text/csv'
	return response

def upload_file_currenttweets_hashtags(request):
	if request.method == 'POST':
		f= pd.read_csv(request.FILES["document"])
		#df=pd.DataFrame(f,columns=["username","items"])
		df=pd.DataFrame(f)

		data = pd.DataFrame([])
	
		for index, row in df.head().iterrows():
			auth = tweepy.OAuthHandler(CONSUMER_KEY, CONSUMER_SECRET)
			auth.set_access_token(ACCESS_TOKEN, ACCESS_TOKEN_SECRET)
			api = tweepy.API(auth,wait_on_rate_limit=True)
  
			#users= tweepy.Cursor(api.user_timeline,screen_name=row['username'],count=200).items(int(row['items']))
			#users= tweepy.Cursor(api.search, q= row['username'],count=200).items(int(row['items']))
			users= tweepy.Cursor(api.search, q='#' + row['username'],count=200).items(int(row['items']))
			outtweets=[]
			outtweets=[[tweet.id_str,tweet.user.verified,tweet.user.created_at,tweet.user.screen_name,tweet.user.name.encode('utf-8'),tweet.user.description.encode('utf-8'),tweet.user.statuses_count,tweet.user.friends_count,tweet.user.followers_count,tweet.created_at,tweet.text.encode('utf-8'),len(tweet.text),tweet.source,tweet.favorite_count,tweet.retweet_count,tweet.user.location.encode('utf-8'),tweet.entities.get('media', [])] for tweet in users]
			#outtweets=[[tweet.id_str,tweet.user.screen_name,tweet.user.name.encode('utf-8'),tweet.user.description.encode('utf-8'),tweet.user.statuses_count,tweet.user.followers_count,tweet.created_at,tweet.text.encode('utf-8'),len(tweet.text),tweet.source,tweet.favorite_count,tweet.retweet_count,tweet.user.location.encode('utf-8')] for tweet in users]
			#outtweets=[[tweet.id_str,tweet.user.screen_name,tweet.created_at,tweet.text.encode('utf-8'),len(tweet.text),tweet.source,tweet.favorite_count,tweet.retweet_count,tweet.user.location] for tweet in users]
			#outtweets = [[d.id,d.created_at,d.user.location,d.text] for d in users]
			outtweets= pd.DataFrame(outtweets)
			#outtweets.columns=['id','created_at','location','text']
			#outtweets.columns=['id','screenname','created_at','text','length','source','favorite_count','retweets','location']
			outtweets.columns=['id','verified','user_created_at','screenname','name','description','Tweets_count','friends_count','follower_count','created_at','text','length','source','favorite_count','retweets','location','media_url']
			#outtweets.columns=['id','screenname','name','description','Tweets_count','follower_count','created_at','text','length','source','favorite_count','retweets','location']
			data=data.append(outtweets)

			#stream.seek(0)
	if request.method=='POST':
		
		form = EmailForm(request.POST)
		if form.is_valid():
			recipient_email_address=form.cleaned_data["recipient_email_address"]
			email = EmailMessage("File Downloaded", "Your file has been downloaded", "ILE@rennes-sb.com", [recipient_email_address])
			email.send(fail_silently=False)		
		response = HttpResponse(data.to_csv(),email)
		
		response['Content-Disposition'] = "attachment; filename=userscurrenttweets.csv"
		response['Content-Type'] = 'text/csv'
		return response
	else:
		return HttpResponse("please enter proper data file")

def oldtweets(request):
	if request.method=='POST':
		form=OldTweetForm(request.POST)
		if form.is_valid():
			username = form.cleaned_data["username"]
			since=form.cleaned_data["since"]
			until=form.cleaned_data["until"]
			maxtweets = form.cleaned_data["maxtweets"]
			recipient_email_address=form.cleaned_data["recipient_email_address"]
			Criteria = got3.manager.TweetCriteria().setUsername(username).setSince(datetime.strftime(since,'%Y-%m-%d')).setUntil(datetime.strftime(until,'%Y-%m-%d')).setMaxTweets(int(maxtweets))
			tweets = got3.manager.TweetManager.getTweets(Criteria)
			df = pd.DataFrame(data=[tweet.text.encode('utf-8') for tweet in tweets], columns=['Tweets'])
			df['id'] = np.array([tweet.id for tweet in tweets])
			#df['len'] = np.array([len(tweet.text) for tweet in tweets])
			df['username'] = np.array([tweet.username for tweet in tweets])
			df['date'] = np.array([tweet.date for tweet in tweets])
			df['favorites'] = np.array([tweet.favorites for tweet in tweets])
			df['retweets'] = np.array([tweet.retweets for tweet in tweets])
			df['hashtags'] = np.array([tweet.hashtags for tweet in tweets])
			df['mentions'] = np.array([tweet.mentions for tweet in tweets])
			df['geo']=np.array([tweet.geo for tweet in tweets])
			
			#print(df)
			#df.to_csv("Twitterthousanddat12e.csv", index=False)
			email = EmailMessage("File Downloaded", "Your file has been downloaded", "ILE@rennes-sb.com", [recipient_email_address])
			email.send(fail_silently=False)
			response = HttpResponse(df.to_csv(),email)
			response['Content-Disposition'] = "attachment; filename=oldtweets.csv"
			response['Content-Type'] = 'text/csv'
			return response
			

	else:
		form=OldTweetForm()

	return render(request,'oldtweets.html',{})




def upload_file_oldtweets(request):
	if request.method == 'POST':
		f= pd.read_csv(request.FILES["document"])
		#df=pd.DataFrame(f,columns=["username","items"])
		#df=pd.DataFrame(f)
		df=pd.DataFrame(f)
		print(df)
		#df.save()
		data = pd.DataFrame([])
	
		for index, row in df.head().iterrows():
		
			Criteria = got3.manager.TweetCriteria().setUsername(row['username']).setSince(row['since']).setUntil(row['until']).setMaxTweets(int(row['maxtweets']))
			tweets = got3.manager.TweetManager.getTweets(Criteria)
			df = pd.DataFrame(data=[tweet.text.encode('utf-8') for tweet in tweets], columns=['Tweets'])
			df['id'] = np.array([tweet.id for tweet in tweets])
			df['len'] = np.array([len(tweet.text) for tweet in tweets])
			df['username'] = np.array([tweet.username for tweet in tweets])
			df['date'] = np.array([tweet.date for tweet in tweets])
			df['favorites'] = np.array([tweet.favorites for tweet in tweets])
			df['retweets'] = np.array([tweet.retweets for tweet in tweets])
			df['hashtags'] = np.array([tweet.hashtags for tweet in tweets])
			df['mentions'] = np.array([tweet.mentions for tweet in tweets])
			df['geo']=np.array([tweet.geo for tweet in tweets])
			#df3=df.append()
			#df2= pd.merge(df,df)
			#df3=df.append('df')
			#df3=pd.concat(df)
			data=data.append(df)
			print(data)
	if request.method=='POST':
		
		form = EmailForm(request.POST)
		if form.is_valid():
			recipient_email_address=form.cleaned_data["recipient_email_address"]
			email = EmailMessage("File Downloaded", "Your file has been downloaded", "ILE@rennes-sb.com", [recipient_email_address])
			email.send(fail_silently=False)	
		response = HttpResponse(data.to_csv(),email)
		response['Content-Disposition'] = "attachment; filename=usersoldtweets.csv"
		response['Content-Type'] = 'text/csv'
		return response
	else:
		return HttpResponse("please enter proper data file")
		


def usertimeline(request):
	if request.method=='POST':
		form=GetUserTimelineForm(request.POST)
		if form.is_valid():
			username=form.cleaned_data['username']
			items=form.cleaned_data['items']
			recipient_email_address=form.cleaned_data["recipient_email_address"]
			#username = request.POST.get('username')
			#items = request.POST.get('items')
			
			auth = tweepy.OAuthHandler(CONSUMER_KEY, CONSUMER_SECRET)
			auth.set_access_token(ACCESS_TOKEN, ACCESS_TOKEN_SECRET)
			api = tweepy.API(auth,wait_on_rate_limit=True)
  
			users= tweepy.Cursor(api.user_timeline,screen_name=username,count=200).items(int(items))
			outtweets=[]
			outtweets=[[tweet.id_str,tweet.user.screen_name,tweet.created_at,tweet.text.encode('utf-8'),len(tweet.text),tweet.source,tweet.favorite_count,tweet.retweet_count,tweet.user.location] for tweet in users]
			#outtweets = [[d.id,d.created_at,d.user.location,d.text] for d in users]
			outtweets= pd.DataFrame(outtweets)
			#outtweets.columns=['id','created_at','location','text']
			outtweets.columns=['id','screenname','created_at','text','length','source','favorite_count','retweets','location']
			#data=outtweets.to_csv()
			#print(data)
			response = HttpResponse(outtweets.to_csv())
			response['Content-Disposition'] = "attachment; filename=userstimelinetweets.csv"
			response['Content-Type'] = 'text/csv'
			return response
	else:
		form=GetUserTimelineForm()

	return render(request,'getusertimeline.html',{})


def getfollowerlist(request):
	if request.method=='POST':
		form=GetFollowerslistForm(request.POST)
		if form.is_valid():
			username=form.cleaned_data['username']
			items=form.cleaned_data['items']
			recipient_email_address=form.cleaned_data["recipient_email_address"]
			auth = tweepy.OAuthHandler(CONSUMER_KEY, CONSUMER_SECRET)
			auth.set_access_token(ACCESS_TOKEN, ACCESS_TOKEN_SECRET)
			api = tweepy.API(auth,wait_on_rate_limit=True)
			#%%
			users= tweepy.Cursor(api.followers,screen_name=username,count=200).items(int(items))
		

			outtweets = [[d.id,d.name.encode('utf-8'),d.screen_name.encode('utf-8'),d.created_at,d.location] for d in users]
			outtweets= pd.DataFrame(outtweets)
			outtweets.columns=['id','name','screen_name','created_at','location']

			response = HttpResponse(outtweets.to_csv())
			response['Content-Disposition'] = "attachment; filename=followers.csv"
			response['Content-Type'] = 'text/csv'
			return response
			
			
			
	else:
		form=GetFollowerslistForm()

	return render(request,'getusertimeline.html',{})


def upload(request):
	context = {}
	if request.method=="POST":
		uploaded_file=request.FILES['document']
		#print(uploaded_file.name)
		#print(uploaded_file.size)
		fs=FileSystemStorage()
		name=fs.save(uploaded_file.name,uploaded_file)
		context['url']=fs.url(name)
	return render(request,'upload.html',context)


def upload_file_speech(request):
	if request.method == 'POST':
		audio= sr.AudioFile(request.FILES["document"])
		r = sr.Recognizer()
		with audio as source:
			rec=r.record(source)
		print(r.recognize_google(rec))

		type(r.recognize_google(rec))

		text = r.recognize_google(rec)
		document = Document()

		document.add_heading("Transcipted Audio File", 0)
		document.add_paragraph(text)
		
		response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
		response['Content-Disposition'] = 'attachment; filename=download.docx'
		document.save(response)

		return response



	
	
	

def csv_file_upload(request):
	if request.method == 'POST':
		data= pd.read_csv(request.FILES["document"])
		#df=pd.DataFrame(f)
		document=[]
		a=[]
		for i in range(len(data)):
			a=data.iloc[i][0]
			document.append(a)
		
		document[0:5]

		print(document)
		

	doc_clean = [clean(doc).split() for doc in document]

	dictionary = corpora.Dictionary(doc_clean)
	print(dictionary)
	doc_term_matrix = [dictionary.doc2bow(doc) for doc in doc_clean]
	
	if request.method=='POST':
		
		form = lda_mod(request.POST)
		if form.is_valid():
			total_topics=form.cleaned_data["total_topics"]
			number_words = form.cleaned_data["number_words"]
			Lda = gensim.models.ldamodel.LdaModel
			lda_model = Lda(doc_term_matrix, num_topics=int(total_topics), id2word = dictionary, passes=50)
			lda_model.show_topics(num_topics=int(total_topics), num_words=int(number_words))

			topics = lda_model.show_topics(formatted=False)
			data_flat = [w for w_list in doc_clean for w in w_list]
			counter = Counter(data_flat)

			out = []
			for i, topic in topics:
				for word, weight in topic:
					out.append([word, i , weight, counter[word]])

			df_imp_wcount = pd.DataFrame(out, columns=['word', 'topic_id', 'importance', 'word_count'])
			#df_imp_wcount.to_csv("Important word count.csv",encoding = "utf-8", index = False)
			print(df_imp_wcount)
			#return df_imp_wcount.to_html()
			response = HttpResponse(df_imp_wcount.to_csv())
			response['Content-Disposition'] = "attachment; filename=csvuploadlda.csv"
			response['Content-Type'] = 'text/csv'
			return response
	
		



def removeUnicode(text):
	""" Removes unicode strings like "\u002c" and "x96" """
	text = re.sub(r'(\\u[0-9A-Fa-f]+)',r'', text)       
	text = re.sub(r'[^\x00-\x7f]',r'',text)
	return text

def replaceURL(text):
	""" Replaces url address with "url" """
	text = re.sub('((www\.[^\s]+)|(https?://[^\s]+))','url',text)
	text = re.sub(r'#([^\s]+)', r'\1', text)
	return text

def removeNumbers(text):
	""" Removes integers """
	text = ''.join([i for i in text if not i.isdigit()])         
	return text

def replaceMultiExclamationMark(text):
	""" Replaces repetitions of exlamation marks """
	text = re.sub(r"(\!)\1+", ' multiExclamation ', text)
	return text

def replaceMultiQuestionMark(text):
	""" Replaces repetitions of question marks """
	text = re.sub(r"(\?)\1+", ' multiQuestion ', text)
	return text

def replaceMultiStopMark(text):
	""" Replaces repetitions of stop marks """
	text = re.sub(r"(\.)\1+", ' multiStop ', text)
	return text

def countMultiExclamationMarks(text):
	""" Replaces repetitions of exlamation marks """
	return len(re.findall(r"(\!)\1+", text))

def countMultiQuestionMarks(text):
	""" Count repetitions of question marks """
	return len(re.findall(r"(\?)\1+", text))
	
def countMultiStopMarks(text):
	""" Count repetitions of stop marks """
	return len(re.findall(r"(\.)\1+", text))

def countElongated(text):
	""" Input: a text, Output: how many words are elongated """
	regex = re.compile(r"(.)\1{2}")
	return len([word for word in text.split() if regex.search(word)])

def countAllCaps(text):
	""" Input: a text, Output: how many words are all caps """
	return len(re.findall("[A-Z0-9]{3,}", text))

def countAllCaps(text):
	""" Input: a text, Output: how many words are all caps """
	return len(re.findall("[A-Z0-9]{3,}", text))
def remove_url(txt):
	"""Replace URLs found in a text string with nothing 
	(i.e. it will remove the URL from the string).

	Parameters
	----------
	txt : string
		A text string that you want to parse and remove urls.

	Returns
	-------
	The same txt string with url's removed.
	"""

	return " ".join(re.sub("([^0-9A-Za-z \t])|(\w+:\/\/\S+)", "", txt).split())

def removeEmoticons(text):
	""" Removes emoticons from text """
	text = re.sub(':\)|;\)|:-\)|\(-:|:-D|=D|:P|xD|X-p|\^\^|:-*|\^\.\^|\^\-\^|\^\_\^|\,-\)|\)-:|:\'\(|:\(|:-\(|:\S|T\.T|\.\_\.|:<|:-\S|:-<|\*\-\*|:O|=O|=\-O|O\.o|XO|O\_O|:-\@|=/|:/|X\-\(|>\.<|>=\(|D:', '', text)
	return text

def removeMultiExclamationMark(text):
	""" Replaces repetitions of exlamation marks """
	text = re.sub(r"(\!)\1+", ' ', text)
	return text

def removeMultiQuestionMark(text):
	""" Replaces repetitions of question marks """
	text = re.sub(r"(\?)\1+", ' ', text)
	return text

def removeMultiStopMark(text):
	""" Replaces repetitions of stop marks """
	text = re.sub(r"(\.)\1+", ' ', text)
	return text


def remove_hash(txt):
	"""Removes Hash"""
	txt = re.sub(r'#', '', txt)
	return txt

def countEmoticons(text):
	""" Input: a text, Output: how many emoticons """
	return len(re.findall(':\)|;\)|:-\)|\(-:|:-D|=D|:P|xD|X-p|\^\^|:-*|\^\.\^|\^\-\^|\^\_\^|\,-\)|\)-:|:\'\(|:\(|:-\(|:\S|T\.T|\.\_\.|:<|:-\S|:-<|\*\-\*|:O|=O|=\-O|O\.o|XO|O\_O|:-\@|=/|:/|X\-\(|>\.<|>=\(|D:', text))

def remove_at(txt):
	"""remove @mentions"""
	txt = re.sub(r'@[A-Za-z0-9]+', '', txt)
	return txt


#######  PreProcessing  #########
stop = set(stopwords.words('english'))
exclude = set(string.punctuation) 
lemma = WordNetLemmatizer()

def clean(doc):
	stop_free = " ".join([i for i in doc.lower().split() if i not in stop])
	punc_free = ''.join(ch for ch in stop_free if ch not in exclude)
	normalized = " ".join(lemma.lemmatize(word) for word in punc_free.split())
	unicode = removeUnicode(normalized)
	rmnumbers = removeNumbers(unicode)
	multiexclam = replaceMultiExclamationMark(rmnumbers)
	multiques = replaceMultiQuestionMark(multiexclam)
	multistop = replaceMultiStopMark(multiques)
	return multistop







from gensim import models
def topics_document_to_dataframe(topics_document, num_topics):
	res = pd.DataFrame(columns=range(num_topics))
	for topic_weight in topics_document:
		res.loc[0,topic_weight[0]] =topic_weight[1]
	return res


def format_sentence(request):
	if request.method == 'POST':
		data= pd.read_csv(request.FILES["document"])
		#df=pd.DataFrame(f)
		document=[]
		a=[]
		for i in range(len(data)):
			a=data.iloc[i][0]
			document.append(a)
		
		document[0:5]
	
	doc_clean = [clean(doc).split() for doc in document]

	dictionary = corpora.Dictionary(doc_clean)
	print(dictionary)
	doc_term_matrix = [dictionary.doc2bow(doc) for doc in doc_clean]
	

	
	if request.method=='POST':
		
		form = csv_lda_mod(request.POST)
		if form.is_valid():

	
		
			total_topics=form.cleaned_data["total_topics"]
			number_words =form.cleaned_data["number_words"]
			recipient_email_address=form.cleaned_data["recipient_email_address"]
			#user = str(get(user.email))
			# column names
			topicnames = ["Topic" + str(i) for i in range(total_topics)]
			# index names
			docnames = ["Doc" + str(i) for i in range(len(data))]

			tfidf = models.TfidfModel(doc_term_matrix)
			corpus_tfidf = tfidf[doc_term_matrix]
			
			Lda = gensim.models.ldamodel.LdaModel
			lda_model = Lda(corpus_tfidf, num_topics=int(total_topics), id2word = dictionary, passes=50)
			lda_model.show_topics(num_topics=int(total_topics), num_words=int(number_words))
			topics_1 = [lda_model[doc_term_matrix[i]] for i in range(len(data))]
			


			topics = lda_model.show_topics(num_topics=int(total_topics), num_words=int(number_words),formatted=False)
			data_flat = [w for w_list in doc_clean for w in w_list]
			counter = Counter(data_flat)

			# Init output
			sent_topics_df = pd.DataFrame([])

			# Get main topic in each document
			for i, row_list in enumerate(lda_model[doc_term_matrix]):
				row = row_list[0] if lda_model.per_word_topics else row_list            
				# print(row)
				row = sorted(row, key=lambda x: (x[1]), reverse=True)
				# Get the Dominant topic, Perc Contribution and Keywords for each document
				for j, (topic_num, prop_topic) in enumerate(row):
					if j == 0:  # => dominant topic
						wp = lda_model.show_topic(topic_num)
						topic_keywords = ", ".join([word for word, prop in wp])
						sent_topics_df = sent_topics_df.append(pd.Series([int(topic_num), round(prop_topic,4), topic_keywords]), ignore_index=True)
					else:
						break
			sent_topics_df.columns = ['Dominant_Topic', 'Perc_Contribution', 'Topic_Keywords']

			# Add original text to the end of the output
			contents = pd.Series(document)
			sent_topics_df = pd.concat([sent_topics_df, contents], axis=1)
			#return(sent_topics_df)


			#df_topic_sents_keywords = format_topics_sentences(lda_model,doc_term_matrix,doc_clean)

			# Format
			df_dominant_topic = sent_topics_df.reset_index()
			df_dominant_topic.columns = ['Document_No', 'Dominant_Topic', 'Topic_Perc_Contrib', 'Keywords', 'Text']
			print(df_dominant_topic.head(10))
			out = []
			for i, topic in topics:
				for word, weight in topic:
					out.append([word, i , weight, counter[word]])

			df_imp_wcount = pd.DataFrame(out, columns=['word', 'topic_id', 'importance', 'word_count'])
			#df_imp_wcount.to_csv("Important word count.csv",encoding = "utf-8", index = False)
			print(df_imp_wcount)
			# Like TF-IDF, create a matrix of topic weighting, with documents as rows and topics as columns
			document_topic = pd.concat([topics_document_to_dataframe(topics_document, num_topics=int(total_topics)) for topics_document in topics_1]).reset_index(drop=True).fillna(0)
			document_topic.columns = topicnames
			#document_topic.to_csv("Topic Distribution.csv",encoding = "utf-8" ,index = False)
			#document_topic.to_csv("Topic Distribution.csv",encoding = "utf-8" ,index = False)
			# Word Frequencies before and after preprocessing
			# Before pre-processing

			document_1 = " ".join(document)  
			document_1_list = document_1.split()
			document_1_unique = set(document_1_list)
			document_1_dict = {}
			for words in document_1_unique:
				document_1_dict[words] = document_1_list.count(words)
			unclean_words = pd.Series(list(document_1_dict.keys()))
			unclean_freq = pd.Series(list(document_1_dict.values()))
			unclean_word_df = pd.DataFrame()
			unclean_word_df = pd.concat([unclean_words,unclean_freq],axis = 1)
			unclean_word_df.columns = ["Words Before Preprocessing","Frequencies"]
			unclean_word_df = unclean_word_df.sort_values(by=["Words Before Preprocessing"])
			#unclean_word_df.to_csv("Unclean Word Frequencies.csv", encoding = "utf-8", index = False)

			docz = [" ".join(doc) for doc in doc_clean]
			docz_1 = " ".join(docz)
			docz_1_list = docz_1.split()
			docz_1_unique = set(docz_1_list)
			docz_1_dict = {}
			for word in docz_1_unique:
				docz_1_dict[word] = docz_1_list.count(word)
			clean_words = pd.Series(list(docz_1_dict.keys()))
			clean_freq = pd.Series(list(docz_1_dict.values()))
			clean_word_df = pd.DataFrame()
			clean_word_df = pd.concat([clean_words,clean_freq],axis = 1)
			clean_word_df.columns = ["Words After Preprocessing","Frequencies"]
			clean_word_df = clean_word_df.sort_values(by=["Words After Preprocessing"])
			#clean_word_df.to_csv("Clean Word Frequencies.csv", encoding = "utf-8", index = False)
			Topic_Modeling = pd.concat([df_imp_wcount, df_dominant_topic,document_topic,unclean_word_df,clean_word_df], axis=1,sort=False)
			#response = HttpResponse(Topic_Modeling.to_csv())

			#df_dominant_topic.to_csv()
			#return df_dominant_topic.to_html()
			#print(user_email)
			email = EmailMessage("File Downloaded", "Your file has been downloaded", "ILE@rennes-sb.com", [recipient_email_address])
			#attachment = MIMENonMultipart('text', 'csv', charset='utf-8')
			#attachment.add_header("Content-Disposition", "attachment", filename=df_dominant_topic.csv)
			#email.attach(attachment)
			email.send(fail_silently=False)
			response = HttpResponse(Topic_Modeling.to_csv(),email)
			#response = HttpResponse(df_dominant_topic.to_csv(),send_mail('Subject here','Here is the message.','ramumunnangi96@gmail.com',[recipient_email_address,],fail_silently=False,))
			#response = HttpResponse(df_dominant_topic.to_csv(),email)
			
			response['Content-Disposition'] = "attachment; filename=df_dominant_topic.csv"
			response['Content-Type'] = 'text/csv'
			return response
	else:
		form = csv_lda_mod()
	return render(request,'uploadcsv.html',{})


def text_file_upload(request):
	
	if request.method == 'POST':
		files =request.FILES.getlist("document")
		

		#print(files)
		
		data=[]
		for file in files:
			#text.append(str(file))
			df=file.read()
			data.append(str(df))
			
		print(data)
		print(type(data))
		#print(b''.join(data).decode())
		data[0]
	doc_clean = [clean(doc).split() for doc in data]

	dictionary = corpora.Dictionary(doc_clean)
	print(dictionary)
	doc_term_matrix = [dictionary.doc2bow(doc) for doc in doc_clean]
	#return render(request,'preparar_pdf.html',{"lista": lista})
	if request.method=='POST':
		
		form = text_lda_mod(request.POST)
		if form.is_valid():
			total_topics=form.cleaned_data["total_topics"]
			number_words = form.cleaned_data["number_words"]
			Lda = gensim.models.ldamodel.LdaModel
			lda_model = Lda(doc_term_matrix, num_topics=int(total_topics), id2word = dictionary, passes=50)
			lda_model.show_topics(num_topics=int(total_topics), num_words=int(number_words))

			topics = lda_model.show_topics(formatted=False)
			data_flat = [w for w_list in doc_clean for w in w_list]
			counter = Counter(data_flat)

			out = []
			for i, topic in topics:
				for word, weight in topic:
					out.append([word, i , weight, counter[word]])

			df_imp_wcount = pd.DataFrame(out, columns=['word', 'topic_id', 'importance', 'word_count'])
			#df_imp_wcount.to_csv("Important word count.csv",encoding = "utf-8", index = False)
			print(df_imp_wcount)
			#return df_imp_wcount.to_html()
			response = HttpResponse(df_imp_wcount.to_csv())
			response['Content-Disposition'] = "attachment; filename=csvuploadlda.csv"
			response['Content-Type'] = 'text/csv'
			return response
	

		

def text_format_sentence(request):
	if request.method == 'POST':
		files =request.FILES.getlist("document")
		data=[]
		for file in files:
			df=file.read()
			data.append(str(df))
			
		print(data)
		print(type(data))
		
		data[0]
	
	doc_clean = [clean(doc).split() for doc in data]

	dictionary = corpora.Dictionary(doc_clean)
	print(dictionary)
	doc_term_matrix = [dictionary.doc2bow(doc) for doc in doc_clean]
	if request.method=='POST':
		
		form = text_lda_mode(request.POST)
		if form.is_valid():
			total_topics=form.cleaned_data["total_topics"]
			number_words =form.cleaned_data["number_words"]
			recipient_email_address=form.cleaned_data["recipient_email_address"]
			#user = str(get(user.email))
			# column names
			topicnames = ["Topic" + str(i) for i in range(total_topics)]
			# index names
			docnames = ["Doc" + str(i) for i in range(len(data))]

			tfidf = models.TfidfModel(doc_term_matrix)
			corpus_tfidf = tfidf[doc_term_matrix]
			
			Lda = gensim.models.ldamodel.LdaModel
			lda_model = Lda(corpus_tfidf, num_topics=int(total_topics), id2word = dictionary, passes=50)
			lda_model.show_topics(num_topics=int(total_topics), num_words=int(number_words))
			topics_1 = [lda_model[doc_term_matrix[i]] for i in range(len(data))]
			


			topics = lda_model.show_topics(num_topics=int(total_topics), num_words=int(number_words),formatted=False)
			data_flat = [w for w_list in doc_clean for w in w_list]
			counter = Counter(data_flat)

			# Init output
			sent_topics_df = pd.DataFrame([])

			# Get main topic in each document
			for i, row_list in enumerate(lda_model[doc_term_matrix]):
				row = row_list[0] if lda_model.per_word_topics else row_list            
				# print(row)
				row = sorted(row, key=lambda x: (x[1]), reverse=True)
				# Get the Dominant topic, Perc Contribution and Keywords for each document
				for j, (topic_num, prop_topic) in enumerate(row):
					if j == 0:  # => dominant topic
						wp = lda_model.show_topic(topic_num)
						topic_keywords = ", ".join([word for word, prop in wp])
						sent_topics_df = sent_topics_df.append(pd.Series([int(topic_num), round(prop_topic,4), topic_keywords]), ignore_index=True)
					else:
						break
			sent_topics_df.columns = ['Dominant_Topic', 'Perc_Contribution', 'Topic_Keywords']

			# Add original text to the end of the output
			contents = pd.Series(document)
			sent_topics_df = pd.concat([sent_topics_df, contents], axis=1)
			#return(sent_topics_df)


			#df_topic_sents_keywords = format_topics_sentences(lda_model,doc_term_matrix,doc_clean)

			# Format
			df_dominant_topic = sent_topics_df.reset_index()
			df_dominant_topic.columns = ['Document_No', 'Dominant_Topic', 'Topic_Perc_Contrib', 'Keywords', 'Text']
			print(df_dominant_topic.head(10))
			out = []
			for i, topic in topics:
				for word, weight in topic:
					out.append([word, i , weight, counter[word]])

			df_imp_wcount = pd.DataFrame(out, columns=['word', 'topic_id', 'importance', 'word_count'])
			#df_imp_wcount.to_csv("Important word count.csv",encoding = "utf-8", index = False)
			print(df_imp_wcount)
			# Like TF-IDF, create a matrix of topic weighting, with documents as rows and topics as columns
			document_topic = pd.concat([topics_document_to_dataframe(topics_document, num_topics=int(total_topics)) for topics_document in topics_1]).reset_index(drop=True).fillna(0)
			document_topic.columns = topicnames
			#document_topic.to_csv("Topic Distribution.csv",encoding = "utf-8" ,index = False)
			#document_topic.to_csv("Topic Distribution.csv",encoding = "utf-8" ,index = False)
			# Word Frequencies before and after preprocessing
			# Before pre-processing

			document_1 = " ".join(data)  
			document_1_list = document_1.split()
			document_1_unique = set(document_1_list)
			document_1_dict = {}
			for words in document_1_unique:
				document_1_dict[words] = document_1_list.count(words)
			unclean_words = pd.Series(list(document_1_dict.keys()))
			unclean_freq = pd.Series(list(document_1_dict.values()))
			unclean_word_df = pd.DataFrame()
			unclean_word_df = pd.concat([unclean_words,unclean_freq],axis = 1)
			unclean_word_df.columns = ["Words Before Preprocessing","Frequencies"]
			unclean_word_df = unclean_word_df.sort_values(by=["Words Before Preprocessing"])
			#unclean_word_df.to_csv("Unclean Word Frequencies.csv", encoding = "utf-8", index = False)

			docz = [" ".join(doc) for doc in doc_clean]
			docz_1 = " ".join(docz)
			docz_1_list = docz_1.split()
			docz_1_unique = set(docz_1_list)
			docz_1_dict = {}
			for word in docz_1_unique:
				docz_1_dict[word] = docz_1_list.count(word)
			clean_words = pd.Series(list(docz_1_dict.keys()))
			clean_freq = pd.Series(list(docz_1_dict.values()))
			clean_word_df = pd.DataFrame()
			clean_word_df = pd.concat([clean_words,clean_freq],axis = 1)
			clean_word_df.columns = ["Words After Preprocessing","Frequencies"]
			clean_word_df = clean_word_df.sort_values(by=["Words After Preprocessing"])
			#clean_word_df.to_csv("Clean Word Frequencies.csv", encoding = "utf-8", index = False)
			Topic_Modeling = pd.concat([df_imp_wcount, df_dominant_topic,document_topic,unclean_word_df,clean_word_df], axis=1,sort=False)
			#response = HttpResponse(Topic_Modeling.to_csv())

			#df_dominant_topic.to_csv()
			#return df_dominant_topic.to_html()
			#print(user_email)
			email = EmailMessage("File Downloaded", "Your file has been downloaded", "ILE@rennes-sb.com", [recipient_email_address])
			#attachment = MIMENonMultipart('text', 'csv', charset='utf-8')
			#attachment.add_header("Content-Disposition", "attachment", filename=df_dominant_topic.csv)
			#email.attach(attachment)
			email.send(fail_silently=False)
			response = HttpResponse(Topic_Modeling.to_csv(),email)
			#response = HttpResponse(df_dominant_topic.to_csv(),send_mail('Subject here','Here is the message.','ramumunnangi96@gmail.com',[recipient_email_address,],fail_silently=False,))
			#response = HttpResponse(df_dominant_topic.to_csv(),email)
			
			response['Content-Disposition'] = "attachment; filename=df_dominant_topic.csv"
			response['Content-Type'] = 'text/csv'
			return response

	
		
			




def clean_data(doc):
	stop_free = " ".join([i for i in doc.lower().split() if i not in stop])
	punc_free = ''.join(ch for ch in stop_free if ch not in exclude)
	normalized = " ".join(lemma.lemmatize(word) for word in punc_free.split())
	urlr = remove_url(normalized)
	unicode = removeUnicode(urlr)
	hasht = remove_hash(unicode)
	at_user = remove_at(hasht)
	rmnumbers = removeNumbers(at_user)
	multiexclam = removeMultiExclamationMark(rmnumbers)
	multiques = removeMultiQuestionMark(multiexclam)
	multistop = removeMultiStopMark(multiques)
	emote = removeEmoticons(multistop)
	return emote

#doc_clean_data = [clean_data(doc) for doc in document]
def get_tweet_sentiment(tweet):
  
	blob = TextBlob(tweet)
 
	# get sentiment
	if blob.sentiment.polarity > 0:
		sentiment = 'positive'
	
	elif blob.sentiment.polarity < 0:
		sentiment = 'negative'
	
	else:
		sentiment = 'neutral'
	
 
	return sentiment
document=[]
def get_processed_tweets(tweets):
  
	processed_tweets = []
	j=0
 
	for tweet in tweets:

		tweet_dict = {}
		tweet_dict['text'] = document[j] 
		tweet_dict['sentiment'] = get_tweet_sentiment(tweet)
		j = j+1
 
		processed_tweets.append(tweet_dict)
 
	return processed_tweets


		
def sentiment_analysis(request):
	
	 
	if request.method == 'POST':
		data= pd.read_csv(request.FILES["document"])
		#document=[]
		a=[]
		for i in range(len(data)):
			a=data.iloc[i][0]
			document.append(a)
		
		document[0:5]

		print(document)
		

	doc_clean_data = [clean_data(doc) for doc in document]
	tweets_with_sentiment = get_processed_tweets(doc_clean_data)
	positive_tweets = [tweet for tweet in tweets_with_sentiment if tweet['sentiment'] == 'positive']
	negative_tweets = [tweet for tweet in tweets_with_sentiment if tweet['sentiment'] == 'negative']
	neutral_tweets = [tweet for tweet in tweets_with_sentiment if tweet['sentiment'] == 'neutral']

	positive_percent = 100 * len(positive_tweets) / len(tweets_with_sentiment)
	negative_percent = 100 * len(negative_tweets) / len(tweets_with_sentiment)
	neutral_percent  = 100 * len(neutral_tweets)  / len(tweets_with_sentiment)
 
	print ('Positive Tweets  | Count: {} , Percent: {} % '. format(len(positive_tweets), positive_percent))
	print ('Negative Tweets | Count: {} , Percent: {} %' . format(len(negative_tweets), negative_percent))
	print ('Neutral Tweets  | Count: {} , Percent: {} %' . format(len(neutral_tweets), neutral_percent))
	positive_percent = 100 * len(positive_tweets) / len(tweets_with_sentiment)
	negative_percent = 100 * len(negative_tweets) / len(tweets_with_sentiment)
	neutral_percent  = 100 * len(neutral_tweets)  / len(tweets_with_sentiment)
 
	print ('Positive Tweets  | Count: {} , Percent: {} % '. format(len(positive_tweets), positive_percent))
	print ('Negative Tweets | Count: {} , Percent: {} %' . format(len(negative_tweets), negative_percent))
	print ('Neutral Tweets  | Count: {} , Percent: {} %' . format(len(neutral_tweets), neutral_percent))
	
		#df=pd.DataFrame(data)
		#df['polarity'] = df.apply(lambda x: TextBlob(x['text']).sentiment.polarity, axis=1)
		#df['Polarity_Status'] = ["Positive" if x >0 else "neutral" for x in df['polarity']]
		#df["Polarity_Status"]
		#for x in df['polarity']:
			#if df["polarity"]>0:
				#df["Polarity_Status"]="Positive"
			#elif df["polarity"]<0:
				#df["Polarity_Status"]="Negative"
				
			#else:
				#df["Polarity_Status"]="Neutral"

		


		
   
		
		#df['subjectivity'] = df.apply(lambda x: TextBlob(x['text']).sentiment.subjectivity, axis=1)
		


	response=HttpResponse(tweets_with_sentiment)

		
	return response
	
def pdf_lda(request):
	if request.method == 'POST':
		pdfFileObj =request.FILES.getlist("document")
		#pdfFileObj = request.FILES["document"]
		pdfWriter = PyPDF2.PdfFileWriter() 
		text = ''
		data = []
		for file in pdfFileObj:
			#text.append(str(file))
			#df=file.read()
			#data.append(str(df))
			with open(file.temporary_file_path(), 'rb') as f: 
				#pdfMerger.append(f)
				pdfReader = PyPDF2.PdfFileReader(f)
				print(pdfReader.numPages)
				pageObj = pdfReader.getPage(0)
				text=text+pageObj.extractText()
				data.append(text)
		
		
		print(data)
		print(type(data))
	doc_clean = [clean(doc).split() for doc in data]

	dictionary = corpora.Dictionary(doc_clean)
	print(dictionary)
	doc_term_matrix = [dictionary.doc2bow(doc) for doc in doc_clean]
	#return render(request,'preparar_pdf.html',{"lista": lista})
	if request.method=='POST':
		
		form = pdf_lda_mod(request.POST)
		if form.is_valid():
			total_topics=form.cleaned_data["total_topics"]
			number_words = form.cleaned_data["number_words"]
			Lda = gensim.models.ldamodel.LdaModel
			lda_model = Lda(doc_term_matrix, num_topics=int(total_topics), id2word = dictionary, passes=50)
			lda_model.show_topics(num_topics=int(total_topics), num_words=int(number_words))

			topics = lda_model.show_topics(formatted=False)
			data_flat = [w for w_list in doc_clean for w in w_list]
			counter = Counter(data_flat)

			out = []
			for i, topic in topics:
				for word, weight in topic:
					out.append([word, i , weight, counter[word]])

			df_imp_wcount = pd.DataFrame(out, columns=['word', 'topic_id', 'importance', 'word_count'])
			#df_imp_wcount.to_csv("Important word count.csv",encoding = "utf-8", index = False)
			print(df_imp_wcount)
			#return df_imp_wcount.to_html()
			response = HttpResponse(df_imp_wcount.to_csv())
			response['Content-Disposition'] = "attachment; filename=csvuploadlda.csv"
			response['Content-Type'] = 'text/csv'
			return response
				
				

		
		
def pdf_dominant_topic_lda(request):
	if request.method == 'POST':
		pdfFileObj =request.FILES.getlist("document")
		#pdfFileObj = request.FILES["document"]
		pdfWriter = PyPDF2.PdfFileWriter() 
		text = ''
		data = []
		for file in pdfFileObj:
			#text.append(str(file))
			#df=file.read()
			#data.append(str(df))
			with open(file.temporary_file_path(), 'rb') as f: 
				#pdfMerger.append(f)
				pdfReader = PyPDF2.PdfFileReader(f)
				print(pdfReader.numPages)
				pageObj = pdfReader.getPage(0)
				text=text+pageObj.extractText()
				data.append(text)
		
		
		print(data)
		print(type(data))
	doc_clean = [clean(doc).split() for doc in data]

	dictionary = corpora.Dictionary(doc_clean)
	print(dictionary)
	doc_term_matrix = [dictionary.doc2bow(doc) for doc in doc_clean]
	if request.method=='POST':
		
		form = pdf_lda_mode(request.POST)
		if form.is_valid():
			total_topics=form.cleaned_data["total_topics"]
			number_words =form.cleaned_data["number_words"]
			recipient_email_address=form.cleaned_data["recipient_email_address"]
			#user = str(get(user.email))
			# column names
			topicnames = ["Topic" + str(i) for i in range(total_topics)]
			# index names
			docnames = ["Doc" + str(i) for i in range(len(data))]

			tfidf = models.TfidfModel(doc_term_matrix)
			corpus_tfidf = tfidf[doc_term_matrix]
			
			Lda = gensim.models.ldamodel.LdaModel
			lda_model = Lda(corpus_tfidf, num_topics=int(total_topics), id2word = dictionary, passes=50)
			lda_model.show_topics(num_topics=int(total_topics), num_words=int(number_words))
			topics_1 = [lda_model[doc_term_matrix[i]] for i in range(len(data))]
			


			topics = lda_model.show_topics(num_topics=int(total_topics), num_words=int(number_words),formatted=False)
			data_flat = [w for w_list in doc_clean for w in w_list]
			counter = Counter(data_flat)

			# Init output
			sent_topics_df = pd.DataFrame([])

			# Get main topic in each document
			for i, row_list in enumerate(lda_model[doc_term_matrix]):
				row = row_list[0] if lda_model.per_word_topics else row_list            
				# print(row)
				row = sorted(row, key=lambda x: (x[1]), reverse=True)
				# Get the Dominant topic, Perc Contribution and Keywords for each document
				for j, (topic_num, prop_topic) in enumerate(row):
					if j == 0:  # => dominant topic
						wp = lda_model.show_topic(topic_num)
						topic_keywords = ", ".join([word for word, prop in wp])
						sent_topics_df = sent_topics_df.append(pd.Series([int(topic_num), round(prop_topic,4), topic_keywords]), ignore_index=True)
					else:
						break
			sent_topics_df.columns = ['Dominant_Topic', 'Perc_Contribution', 'Topic_Keywords']

			# Add original text to the end of the output
			contents = pd.Series(document)
			sent_topics_df = pd.concat([sent_topics_df, contents], axis=1)
			#return(sent_topics_df)


			#df_topic_sents_keywords = format_topics_sentences(lda_model,doc_term_matrix,doc_clean)

			# Format
			df_dominant_topic = sent_topics_df.reset_index()
			df_dominant_topic.columns = ['Document_No', 'Dominant_Topic', 'Topic_Perc_Contrib', 'Keywords', 'Text']
			print(df_dominant_topic.head(10))
			out = []
			for i, topic in topics:
				for word, weight in topic:
					out.append([word, i , weight, counter[word]])

			df_imp_wcount = pd.DataFrame(out, columns=['word', 'topic_id', 'importance', 'word_count'])
			#df_imp_wcount.to_csv("Important word count.csv",encoding = "utf-8", index = False)
			print(df_imp_wcount)
			# Like TF-IDF, create a matrix of topic weighting, with documents as rows and topics as columns
			document_topic = pd.concat([topics_document_to_dataframe(topics_document, num_topics=int(total_topics)) for topics_document in topics_1]).reset_index(drop=True).fillna(0)
			document_topic.columns = topicnames
			#document_topic.to_csv("Topic Distribution.csv",encoding = "utf-8" ,index = False)
			#document_topic.to_csv("Topic Distribution.csv",encoding = "utf-8" ,index = False)
			# Word Frequencies before and after preprocessing
			# Before pre-processing

			document_1 = " ".join(data)  
			document_1_list = document_1.split()
			document_1_unique = set(document_1_list)
			document_1_dict = {}
			for words in document_1_unique:
				document_1_dict[words] = document_1_list.count(words)
			unclean_words = pd.Series(list(document_1_dict.keys()))
			unclean_freq = pd.Series(list(document_1_dict.values()))
			unclean_word_df = pd.DataFrame()
			unclean_word_df = pd.concat([unclean_words,unclean_freq],axis = 1)
			unclean_word_df.columns = ["Words Before Preprocessing","Frequencies"]
			unclean_word_df = unclean_word_df.sort_values(by=["Words Before Preprocessing"])
			#unclean_word_df.to_csv("Unclean Word Frequencies.csv", encoding = "utf-8", index = False)

			docz = [" ".join(doc) for doc in doc_clean]
			docz_1 = " ".join(docz)
			docz_1_list = docz_1.split()
			docz_1_unique = set(docz_1_list)
			docz_1_dict = {}
			for word in docz_1_unique:
				docz_1_dict[word] = docz_1_list.count(word)
			clean_words = pd.Series(list(docz_1_dict.keys()))
			clean_freq = pd.Series(list(docz_1_dict.values()))
			clean_word_df = pd.DataFrame()
			clean_word_df = pd.concat([clean_words,clean_freq],axis = 1)
			clean_word_df.columns = ["Words After Preprocessing","Frequencies"]
			clean_word_df = clean_word_df.sort_values(by=["Words After Preprocessing"])
			#clean_word_df.to_csv("Clean Word Frequencies.csv", encoding = "utf-8", index = False)
			Topic_Modeling = pd.concat([df_imp_wcount, df_dominant_topic,document_topic,unclean_word_df,clean_word_df], axis=1,sort=False)
			#response = HttpResponse(Topic_Modeling.to_csv())

			#df_dominant_topic.to_csv()
			#return df_dominant_topic.to_html()
			#print(user_email)
			email = EmailMessage("File Downloaded", "Your file has been downloaded", "ILE@rennes-sb.com", [recipient_email_address])
			#attachment = MIMENonMultipart('text', 'csv', charset='utf-8')
			#attachment.add_header("Content-Disposition", "attachment", filename=df_dominant_topic.csv)
			#email.attach(attachment)
			email.send(fail_silently=False)
			response = HttpResponse(Topic_Modeling.to_csv(),email)
			#response = HttpResponse(df_dominant_topic.to_csv(),send_mail('Subject here','Here is the message.','ramumunnangi96@gmail.com',[recipient_email_address,],fail_silently=False,))
			#response = HttpResponse(df_dominant_topic.to_csv(),email)
			
			response['Content-Disposition'] = "attachment; filename=df_dominant_topic.csv"
			response['Content-Type'] = 'text/csv'
			return response

	
		
			

		

	
	
	
	
	
		

		